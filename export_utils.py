# export_utils.py
import io
import pandas as pd
from docx import Document
from docx.shared import Inches

def limpiar_decimales_cero(texto):
    """Elimina el .0 de los strings que representan números enteros perfectos."""
    txt = str(texto)
    if txt.endswith('.0') and txt[:-2].replace('-', '').isdigit():
        return txt[:-2]
    return txt

def generar_word_dinamico(df_data, elementos_renderizados, config):
    doc = Document()
    doc.add_heading(config["app_title"], 0)
    
    doc.add_heading('Resumen de Datos', level=1)
    doc.add_paragraph(f"Registros analizados en esta vista: {len(df_data)}")
    
    for col in config["columnas_filtro"]:
        if col in df_data.columns:
            doc.add_paragraph(f"{col} Diferentes: {df_data[col].nunique()}")
    
    if elementos_renderizados:
        doc.add_heading('Visualizaciones', level=1)
        for titulo, elemento in elementos_renderizados:
            doc.add_heading(titulo, level=2)
            
            if isinstance(elemento, pd.DataFrame):
                elemento_str = elemento.astype(str)
                t = doc.add_table(rows=1, cols=len(elemento_str.columns))
                t.style = 'Table Grid'
                
                # Escribir Encabezados limpios
                for i, col in enumerate(elemento.columns):
                    t.cell(0, i).text = limpiar_decimales_cero(col)
                    
                # Escribir Filas limpias
                for _, row in elemento_str.iterrows():
                    row_cells = t.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = limpiar_decimales_cero(val)
            elif elemento is not None:
                try:
                    img_bytes = elemento.to_image(format="png")
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"[No se pudo exportar el gráfico: {e}]")
            
            doc.add_paragraph()
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
# export_utils.py